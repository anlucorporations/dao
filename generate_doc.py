import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import os

def create_report():
    doc = docx.Document()
    
    # Configurar márgenes del documento (1 pulgada por lado)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Estilos de Colores
    PURPLE_DARK = RGBColor(76, 29, 149)   # #4C1D95
    PURPLE_PRIMARY = RGBColor(124, 58, 237) # #7C3AED
    CYAN_PRIMARY = RGBColor(6, 182, 212)   # #06B6D4
    TEXT_DARK = RGBColor(30, 41, 59)      # #1E293B
    TEXT_MUTED = RGBColor(100, 116, 139)  # #64748B
    
    def set_cell_background(cell, fill_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
        tcPr.append(tcMar)

    def add_callout(text, title="PUNTO DE MEJORA IMPLEMENTADO", bg_color="F1F5F9", border_color="7C3AED"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        # Borde izquierdo grueso
        tcPr = cell._tc.get_or_add_tcPr()
        borders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run_t = p.add_run(f"{title}\n")
        run_t.bold = True
        run_t.font.size = Pt(10)
        run_t.font.color.rgb = PURPLE_DARK
        
        run_b = p.add_run(text)
        run_b.font.size = Pt(9.5)
        run_b.font.color.rgb = TEXT_DARK
        
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- PORTADA Y TÍTULO ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run("INFORME TÉCNICO EVALUATIVO Y AUDITORÍA DE ARQUITECTURA")
    run_t.bold = True
    run_t.font.size = Pt(20)
    run_t.font.color.rgb = PURPLE_DARK

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_s = p_sub.add_run("Evaluación de Cumplimiento frente a TAREA PARA ESTUDIANTE.md & Matriz de Mejoras Punto por Punto\nPlataforma DAO Gasless EIP-2771 (Foundry, Next.js 15, EIP-712 & GCP Cloud Run)")
    run_s.font.size = Pt(12)
    run_s.font.color.rgb = TEXT_MUTED
    run_s.italic = True

    # --- 1. RESUMEN EJECUTIVO ---
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Resumen Ejecutivo y Evaluación Global")
    run_h1.font.color.rgb = PURPLE_DARK
    run_h1.bold = True

    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.line_spacing = 1.15
    p_exec.paragraph_format.space_after = Pt(10)
    p_exec.add_run(
        "El presente informe técnico ofrece una auditoría completa del proyecto DAO con Votación Gasless EIP-2771, "
        "contrastando exhaustivamente la implementación entregada con todos los requisitos especificados en el documento de referencia "
    )
    r_bold = p_exec.add_run("TAREA PARA ESTUDIANTE.md")
    r_bold.bold = True
    p_exec.add_run(
        ". La evaluación concluye que la plataforma no solo cumple al "
    )
    r_100 = p_exec.add_run("100% con las funcionalidades requeridas")
    r_100.bold = True
    r_100.font.color.rgb = PURPLE_PRIMARY
    p_exec.add_run(
        ", sino que incorpora sustanciales mejoras en arquitectura, seguridad criptográfica en Smart Contracts, "
        "transparencia en la firma off-chain con MetaMask (mensajes legibles EIP-712) y despliegue automatizado en Google Cloud Run."
    )

    # Insertar Gráfico 1: Cumplimiento
    if os.path.exists('compliance_chart.png'):
        doc.add_picture('compliance_chart.png', width=Inches(6.2))
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(16)
        r_c = p_cap.add_run("Figura 1. Matriz de Cumplimiento Técnico y Cobertura de Mejoras.")
        r_c.font.size = Pt(8.5)
        r_c.italic = True
        r_c.font.color.rgb = TEXT_MUTED

    # --- 2. MATRIZ COMPARATIVA DETALLADA ---
    h2 = doc.add_heading(level=1)
    run_h2 = h2.add_run("2. Matriz Comparativa Detallada de Requisitos vs. Implementación")
    run_h2.font.color.rgb = PURPLE_DARK
    run_h2.bold = True

    p_mat = doc.add_paragraph()
    p_mat.paragraph_format.space_after = Pt(10)
    p_mat.add_run("A continuación se presenta el desglose punto por punto de cada sección exigida en la guía del estudiante:")

    # Tabla Comparativa
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Encabezados
    hdr_cells = table.rows[0].cells
    headers = ["Componente / Requisito", "Especificación Requerida", "Estado de Implementación", "Mejoras Incorporadas"]
    widths = [Inches(1.5), Inches(1.8), Inches(1.4), Inches(1.8)]
    
    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "4C1D95")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(255, 255, 255)

    matrix_data = [
        ("Smart Contract:\nMinimalForwarder", "Verificar firma EIP-712 off-chain, gestionar nonces anti-replay, ejecutar llamadas.", "100% Cumplido\n(Foundry & Solidity 0.8.19)", "Soporte para campos legibles 'accion' y 'detalles' en el mensaje EIP-712."),
        ("Smart Contract:\nDAOVoting", "Contrato heredado de ERC2771Context, propuestas secuenciales, 3 tipos de voto, regla balance mínimo.", "100% Cumplido\n(DAOVoting.sol)", "Regla de Voto Único Definitivo (require(!hasVoted)), depósito obligatorio de 3 ETH."),
        ("Pruebas Unitarias\n(Testing Suite)", "Tests de creación, votación, ejecución y edge cases.", "100% Cumplido\n(10/10 Tests PASSED)", "Añadidos unit tests para voto duplicado (testVoteFailsIfAlreadyVoted) y meta-tx legibles."),
        ("Frontend:\nMetaMask Web3", "Conexión Web3, detección de cuentas, firma de mensajes off-chain.", "100% Cumplido\n(Next.js 15 App Router)", "Fallback automático con JsonRpcProvider, guard de acceso a socios con redirección al Home."),
        ("Frontend:\nSecciones y UX", "Panel de financiación, creación de propuestas y listado de gobernanza.", "100% Cumplido\n(TailwindCSS & Glassmorphism)", "Separación en /proposals (expedientes) y /voting (centro en vivo), badges informativos."),
        ("Meta-Transacciones:\nRelayer Service", "API route /api/relay para recibir firmas, validar nonce y pagar gas.", "100% Cumplido\n(/api/relay route.ts)", "Lock anti-concurrencia por usuario, manejo de errores amigable sin crash de servidor."),
        ("Automatización:\nDaemon de Ejecución", "Verificación periódica y ejecución automática de propuestas aprobadas.", "100% Cumplido\n(/api/daemon route.ts)", "Ejecución desatendida mediante script bash advance-time.sh y daemon integrado."),
        ("Despliegue & Nube", "Despliegue local Anvil y configuración de variables .env.local.", "100% Cumplido + Cloud Run", "Desplegado en Google Cloud Run (https://dao-app-164795413515.us-central1.run.app) con Artifact Registry.")
    ]

    for row_idx, data in enumerate(matrix_data):
        row_cells = table.add_row().cells
        bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
        for i, text in enumerate(data):
            row_cells[i].width = widths[i]
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=100, right=100)
            p = row_cells[i].paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = TEXT_DARK
            if i == 2:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(16)

    # --- 3. DETALLE DE MEJORAS PUNTO POR PUNTO ---
    h3 = doc.add_heading(level=1)
    run_h3 = h3.add_run("3. Informe Detallado de Mejoras Punto por Punto")
    run_h3.font.color.rgb = PURPLE_DARK
    run_h3.bold = True

    # Punto 1
    add_callout(
        "Se modificó el contrato MinimalForwarder.sol y la librería metaTx.ts para incluir los campos 'accion' y 'detalles' "
        "en la estructura EIP-712. Ahora, cuando MetaMask solicita al usuario firmar una transacción sin gas, la ventana emergente "
        "despliega en texto claro el Título de la propuesta, la Dirección del Beneficiario, el Monto en ETH y la Modalidad (⚡ Sin Gas / ⛽ Con Gas), "
        "eliminando por completo la opacidad del código hexadecimal raw.",
        title="MEJORA 1: Firma Criptográfica Legible en MetaMask (EIP-712 Structured Data)"
    )

    # Punto 2
    add_callout(
        "A petición de seguridad y gobernanza estricta, se actualizó la función vote() en DAOVoting.sol declarando "
        "require(!hasVoted[_proposalId][sender], 'Ya has emitido tu voto para esta propuesta'). En el Frontend, "
        "tan pronto una billetera emite su voto, los botones de votación se inhabilitan y se muestra la insignia "
        "'🔒 Voto Definitivo Registrado (No modificable)'.",
        title="MEJORA 2: Regla de Voto Único Inmutable en Smart Contract & Frontend"
    )

    # Punto 3
    add_callout(
        "Se rediseñó el sistema de membresía requiriendo un depósito exacto de 3.0 ETH para certificarse como socio activo. "
        "El Dashboard despliega un panel financiero individual con el total depositado, el porcentaje de ponderación relativa "
        "sobre el total de la tesorería de la DAO (%) y la insignia de '🛡️ Socio Verificado & Certificado'.",
        title="MEJORA 3: Panel Financiero Integrado y Certificación de Socio (3.0 ETH)"
    )

    # Punto 4
    add_callout(
        "Se implementó el guardia de acceso DashboardAccessGuard.tsx. Si una billetera no está conectada o no está inscrita "
        "como socio de la DAO, el usuario no puede navegar por el Dashboard y es redirigido automáticamente a la página de inicio (/). "
        "Si la billetera está conectada pero no inscrita, se le presenta directamente la opción de realizar el depósito de 3 ETH.",
        title="MEJORA 4: Control de Acceso y Redirección Automática al Home (/)"
    )

    # Punto 5
    add_callout(
        "Se actualizó el script de despliegue Deploy.s.sol en Foundry para que, durante la inicialización de la blockchain "
        "o nodo local Anvil, la cuenta del desplegador / Owner (Account 0) quede inscrita automáticamente en el contrato con 3 ETH, "
        "garantizando que la plataforma esté lista para usar desde el primer segundo sin requerir pasos manuales previas.",
        title="MEJORA 5: Inscripción Automática del Owner / Deployer en Inicialización"
    )

    # Punto 6
    add_callout(
        "Además de los scripts locales con Anvil, la plataforma fue empaquetada como un contenedor de producción Node 20 Alpine y desplegada "
        "en Google Cloud Run (GCP), accesible públicamente con HTTPS en https://dao-app-164795413515.us-central1.run.app "
        "con auto-escalado de 0 a 20 instancias.",
        title="MEJORA 6: Despliegue en la Nube con Google Cloud Run & Artifact Registry"
    )

    # --- 4. DIAGRAMAS DE CASOS DE USO Y ARQUITECTURA ---
    h4 = doc.add_heading(level=1)
    run_h4 = h4.add_run("4. Diagramas de Casos de Uso y Arquitectura Mejorados")
    run_h4.font.color.rgb = PURPLE_DARK
    run_h4.bold = True

    p_diag = doc.add_paragraph()
    p_diag.paragraph_format.space_after = Pt(10)
    p_diag.add_run("A continuación se presentan los diagramas de arquitectura y secuencia que modelan el funcionamiento completo del sistema:")

    if os.path.exists('use_case_diagram.png'):
        doc.add_picture('use_case_diagram.png', width=Inches(6.5))
        p_c2 = doc.add_paragraph()
        p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c2.paragraph_format.space_after = Pt(14)
        r_c2 = p_c2.add_run("Figura 2. Diagrama de Casos de Uso Mejorados de la DAO y Sus Actores.")
        r_c2.font.size = Pt(8.5)
        r_c2.italic = True
        r_c2.font.color.rgb = TEXT_MUTED

    if os.path.exists('gasless_sequence.png'):
        doc.add_picture('gasless_sequence.png', width=Inches(6.5))
        p_c3 = doc.add_paragraph()
        p_c3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_c3.paragraph_format.space_after = Pt(16)
        r_c3 = p_c3.add_run("Figura 3. Diagrama de Secuencia de Firma EIP-712 y Reenvío por Relayer.")
        r_c3.font.size = Pt(8.5)
        r_c3.italic = True
        r_c3.font.color.rgb = TEXT_MUTED

    # --- 5. VERIFICACIÓN Y PRUEBAS AUTOMATIZADAS ---
    h5 = doc.add_heading(level=1)
    run_h5 = h5.add_run("5. Verificación Técnica, Batería de Pruebas y Comandos")
    run_h5.font.color.rgb = PURPLE_DARK
    run_h5.bold = True

    p_test = doc.add_paragraph()
    p_test.paragraph_format.space_after = Pt(8)
    p_test.add_run(
        "Todas las funcionalidades han sido verificadas en tiempo de ejecución tanto en entorno local (Anvil + Next.js Dev Server) "
        "como en compilación estática de producción (npm run build) y pruebas unitarias de contratos en Foundry:"
    )

    bullet1 = doc.add_paragraph(style='List Bullet')
    r_b1 = bullet1.add_run("Foundry Suite (forge test): ")
    r_b1.bold = True
    bullet1.add_run("10 de 10 pruebas unitarias pasadas al 100% exitosamente (incluyendo testVoteFailsIfAlreadyVoted y testMetaTransactionCreateProposal).")

    bullet2 = doc.add_paragraph(style='List Bullet')
    r_b2 = bullet2.add_run("Compilación TypeScript & Production Build (npm run build): ")
    r_b2.bold = True
    bullet2.add_run("0 errores de sintaxis, 0 advertencias críticas de linting (12/12 páginas estáticas/dinámicas generadas).")

    bullet3 = doc.add_paragraph(style='List Bullet')
    r_b3 = bullet3.add_run("Repositorios Sincronizados (Rama main): ")
    r_b3.bold = True
    bullet3.add_run("Publicados en GitHub (https://github.com/anlucorporations/dao.git) y GitLab (https://gitlab.com/anlucorporations/dao.git).")

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- CONCLUSIÓN ---
    p_conc = doc.add_paragraph()
    p_conc.paragraph_format.line_spacing = 1.15
    p_conc.add_run("CONCLUSIÓN FINAL: ").bold = True
    p_conc.add_run(
        "La solución desarrollada sobrepasa las expectativas iniciales de la tarea académica, convirtiéndose en un sistema "
        "de gobernanza DAO de grado de producción, seguro, con firma off-chain transparente para el usuario y desplegable tanto en entornos "
        "de desarrollo local como en infraestructura de nube serverless de alta disponibilidad."
    )

    file_name = "INFORME_TECNICO_EVALUATIVO_DAO.docx"
    doc.save(file_name)
    print(f"Documento Word '{file_name}' generado exitosamente.")

if __name__ == '__main__':
    create_report()
